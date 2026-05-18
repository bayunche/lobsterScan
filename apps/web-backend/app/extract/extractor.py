"""按文件类型抽取文本。返回 (text, status_note)。

text 为空 + status_note 不为空 = 抽取失败 / 不支持,agent 在 data_gaps 里写。
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger("extract")

# 每个文件抽取后最多保留多少字符进 raw_text(避免一份 100 页 PDF 把 prompt 撑爆)
MAX_TEXT_PER_FILE = 6000


@dataclass
class ExtractResult:
    ok: bool
    text: str            # 抽出来的纯文本,失败时空
    note: str            # 状态说明(成功也写,例如 "from .docx, 12 段")
    filename: str
    truncated: bool = False  # 是否因为超过 MAX_TEXT_PER_FILE 被截断


# 后缀白名单 — 走 plain text 路径(任何文本类:文档 / 源代码 / 配置 / 标记)
_TEXT_EXTS: set[str] = {
    # 通用文档
    ".md", ".txt", ".rst", ".org", ".markdown", ".tex",
    # 数据
    ".csv", ".tsv", ".log", ".jsonl", ".ndjson",
    # 配置
    ".yml", ".yaml", ".ini", ".toml", ".cfg", ".conf", ".env", ".properties",
    # 标记
    ".html", ".htm", ".xml", ".svg", ".rss", ".atom",
    # 源代码(常用)
    ".py", ".js", ".ts", ".tsx", ".jsx", ".mjs", ".cjs",
    ".go", ".rs", ".java", ".kt", ".scala", ".swift",
    ".c", ".h", ".cpp", ".hpp", ".cc", ".m", ".mm",
    ".rb", ".php", ".pl", ".lua", ".r",
    ".sh", ".bash", ".zsh", ".fish",
    ".sql", ".graphql", ".proto",
    ".dockerfile", ".containerfile",
    # 笔记
    ".vue", ".svelte", ".astro",
}

# 后缀白名单 — 走 JSON 解析
_JSON_EXTS: set[str] = {".json"}

# 已知二进制类型,直接拒绝(不要 probe 浪费 IO)
_BINARY_MIME_PREFIXES = ("image/", "video/", "audio/", "font/")
_BINARY_EXTS: set[str] = {
    ".doc", ".xls", ".ppt", ".pptx",   # 老 Office
    ".zip", ".tar", ".gz", ".bz2", ".7z", ".rar", ".xz",
    ".exe", ".dll", ".so", ".dylib", ".bin", ".dmg", ".iso",
    ".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".ico", ".heic", ".avif",
    ".mp4", ".mov", ".avi", ".mkv", ".webm", ".flv", ".m4v",
    ".mp3", ".wav", ".flac", ".aac", ".ogg", ".m4a", ".opus",
    ".woff", ".woff2", ".ttf", ".otf", ".eot",
}


def extract_attachment_text(path: Path, filename: str, mime: str) -> ExtractResult:
    """主入口 — 按 mime / 后缀分派,任何文本类内容都尽量抽出来。"""
    if not path.exists():
        return ExtractResult(False, "", "文件已删除/路径不存在", filename)
    ext = path.suffix.lower()
    mime_l = (mime or "").lower()

    try:
        # 1. JSON 走结构化解析(失败回退到 plain)
        if mime_l == "application/json" or ext in _JSON_EXTS:
            r = _read_json(path, filename)
            if r.ok:
                return r
            # JSON 解析失败 → 当 plain text 读

        # 2. Office .docx / .xlsx / .pdf 走专门解析(基于后缀,因为 mime 容易缺)
        if ext == ".docx" or mime_l == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return _read_docx(path, filename)
        if ext == ".xlsx" or mime_l == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return _read_xlsx(path, filename)
        if ext == ".pdf" or mime_l == "application/pdf":
            return _read_pdf(path, filename)

        # 3. Jupyter notebook — .ipynb 是 JSON 结构,抽 cell.source
        if ext == ".ipynb":
            return _read_ipynb(path, filename)

        # 4. 已知二进制类型 → 直接拒绝
        if mime_l.startswith(_BINARY_MIME_PREFIXES) or ext in _BINARY_EXTS:
            return ExtractResult(
                False, "",
                f"{mime_l or ext} 是二进制类型,LLM 看不了。请把要点摘出来贴到文本框",
                filename,
            )

        # 5. mime 是 text/* / application/{xml,javascript,...} → 直接 read
        text_mimes = mime_l.startswith("text/") or mime_l in {
            "application/xml", "application/javascript", "application/typescript",
            "application/x-yaml", "application/yaml", "application/sql",
            "application/x-sh", "application/x-shellscript", "application/x-httpd-php",
            "application/x-python", "application/x-ruby", "application/x-perl",
        }
        if text_mimes or ext in _TEXT_EXTS:
            return _read_plain(path, filename, mime_l or "text/plain")

        # 6. 兜底:**当 UTF-8 文本读一次**,前 4KB 能 decode 无大量 NULL/控制字符就当文本
        return _probe_as_text(path, filename, mime_l, ext)

    except Exception as e:  # noqa: BLE001
        log.warning("extract %s failed: %s", filename, e)
        return ExtractResult(False, "", f"抽取异常: {type(e).__name__}: {str(e)[:120]}", filename)


def _probe_as_text(path: Path, filename: str, mime: str, ext: str) -> ExtractResult:
    """没明确 mime 也不在已知二进制名单时,probe 一下文件头看是不是 UTF-8 文本."""
    try:
        head = path.read_bytes()[:4096]
    except Exception as e:  # noqa: BLE001
        return ExtractResult(False, "", f"读文件头失败: {e}", filename)
    if not head:
        return ExtractResult(True, "", f"空文件(mime={mime} ext={ext})", filename)
    # NULL 字节多 / 控制字符多 → 大概率二进制
    null_ratio = head.count(b"\x00") / len(head)
    if null_ratio > 0.01:
        return ExtractResult(
            False, "",
            f"看起来是二进制文件(NULL 字节占比 {null_ratio:.0%},mime={mime or '未知'} ext={ext}),无法抽文本",
            filename,
        )
    try:
        head.decode("utf-8")
    except UnicodeDecodeError:
        # 试 GBK / Latin-1
        try:
            head.decode("gbk")
            return _read_plain(path, filename, f"{mime or 'text/plain'} (GBK 编码探测)")
        except UnicodeDecodeError:
            return ExtractResult(
                False, "",
                f"非 UTF-8/GBK 文本(mime={mime or '未知'} ext={ext}),无法抽取",
                filename,
            )
    return _read_plain(path, filename, f"{mime or 'text/plain'} (probe-as-text)")


# ─── 具体 extractor ───


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_TEXT_PER_FILE:
        return text, False
    return text[:MAX_TEXT_PER_FILE].rstrip() + "\n…(已截断,共 " + str(len(text)) + " 字)", True


def _read_plain(path: Path, filename: str, mime: str) -> ExtractResult:
    # 先试 UTF-8,失败再 GBK(中文系统老文件常见),最后兜底 errors="replace"
    raw_bytes = path.read_bytes()
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            raw = raw_bytes.decode(enc)
            note = f"from {mime}, {len(raw)} chars"
            if enc != "utf-8":
                note += f" ({enc} 编码)"
            text, trunc = _truncate(raw)
            return ExtractResult(True, text, note, filename, trunc)
        except UnicodeDecodeError:
            continue
    # 兜底:有损解码,加 replace 标记
    raw = raw_bytes.decode("utf-8", errors="replace")
    text, trunc = _truncate(raw)
    return ExtractResult(True, text, f"from {mime}, {len(raw)} chars (有损 utf-8)", filename, trunc)


def _read_ipynb(path: Path, filename: str) -> ExtractResult:
    """Jupyter notebook — 抽 markdown cell 文字 + code cell 源代码 + 关键输出"""
    try:
        j = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:  # noqa: BLE001
        return ExtractResult(False, "", f"ipynb 解析失败: {e}", filename)
    parts: list[str] = []
    cells = j.get("cells", [])
    for i, c in enumerate(cells, 1):
        kind = c.get("cell_type", "?")
        src = c.get("source")
        if isinstance(src, list):
            src = "".join(src)
        src = (src or "").strip()
        if not src:
            continue
        if kind == "markdown":
            parts.append(src)
        elif kind == "code":
            parts.append(f"```python\n{src}\n```")
            # 只摘文本类输出(stream / display_data),跳过图片 base64
            for out in c.get("outputs", [])[:3]:
                ot = out.get("output_type")
                if ot == "stream":
                    txt = out.get("text")
                    if isinstance(txt, list):
                        txt = "".join(txt)
                    if txt:
                        parts.append(f"_out:_ {txt.strip()[:200]}")
                elif ot in ("display_data", "execute_result"):
                    d = out.get("data", {})
                    if "text/plain" in d:
                        tp = d["text/plain"]
                        if isinstance(tp, list):
                            tp = "".join(tp)
                        parts.append(f"_out:_ {tp.strip()[:200]}")
        parts.append("")
    text, trunc = _truncate("\n".join(parts))
    return ExtractResult(True, text, f"from .ipynb, {len(cells)} cells", filename, trunc)


def _read_json(path: Path, filename: str) -> ExtractResult:
    try:
        j = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        return ExtractResult(False, "", f"JSON 解析失败: {e}", filename)
    text, trunc = _truncate(json.dumps(j, ensure_ascii=False, indent=2))
    return ExtractResult(True, text, "from application/json", filename, trunc)


def _read_docx(path: Path, filename: str) -> ExtractResult:
    try:
        from docx import Document
    except ImportError:
        return ExtractResult(False, "", "python-docx 未安装(uv add python-docx)", filename)
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # 表格也抽出来,markdown 风格
    for ti, tbl in enumerate(doc.tables, 1):
        parts.append(f"\n【表 {ti}】")
        for row in tbl.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
        parts.append("")
    text, trunc = _truncate("\n".join(parts))
    return ExtractResult(True, text, f"from .docx, {len(doc.paragraphs)} 段 + {len(doc.tables)} 表", filename, trunc)


def _read_xlsx(path: Path, filename: str) -> ExtractResult:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ExtractResult(False, "", "openpyxl 未安装", filename)
    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    sheet_count = len(wb.sheetnames)
    for sheet_name in wb.sheetnames[:8]:   # 最多前 8 个 sheet
        ws = wb[sheet_name]
        parts.append(f"\n【Sheet: {sheet_name}】")
        # 第一遍扫:确定实际有内容的最大列数(避免 openpyxl 把 trailing 空列也输出)
        max_used_col = 0
        rows_buf: list[list[str]] = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            if row_idx >= 200:               # 单 sheet 最多 200 行,防失控
                parts.append("…(行截断,共 200+ 行)")
                break
            cells = ["" if v is None else str(v).strip().replace("\n", " ") for v in row]
            # 记录这一行最右非空 cell 的位置
            last_nonempty = 0
            for i, c in enumerate(cells):
                if c:
                    last_nonempty = i + 1
            if last_nonempty > max_used_col:
                max_used_col = last_nonempty
            rows_buf.append(cells)
        # 第二遍:裁掉 trailing 空列,只输出有内容的行
        for cells in rows_buf:
            trimmed = cells[:max_used_col]
            if any(trimmed):
                parts.append("| " + " | ".join(trimmed) + " |")
        parts.append("")
    wb.close()
    text, trunc = _truncate("\n".join(parts))
    return ExtractResult(True, text, f"from .xlsx, {sheet_count} sheets", filename, trunc)


def _read_pdf(path: Path, filename: str) -> ExtractResult:
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        return ExtractResult(False, "", "pdfminer.six 未安装", filename)
    raw = extract_text(str(path)) or ""
    # 去掉多余空行
    cleaned = "\n".join(line.rstrip() for line in raw.splitlines() if line.strip())
    text, trunc = _truncate(cleaned)
    return ExtractResult(True, text, f"from .pdf, {len(cleaned)} chars", filename, trunc)
